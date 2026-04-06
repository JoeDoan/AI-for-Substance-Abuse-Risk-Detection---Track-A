from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ── Screenshot paths (adversarial test screenshots from previous session) ─────
SS_DIR = "/Users/dudoan/.gemini/antigravity/brain/f41e723d-db42-4f5f-aa05-9cd4f900184a"
SCREENSHOTS = {
    "home":  f"{SS_DIR}/homepage_before_input_1775155875887.png",
    "full":  f"{SS_DIR}/full_analysis_results_high_risk_1775155920313.png",
    "adv1a": f"{SS_DIR}/adversarial_test1_slippery_slope_part1_1775157003543.png",
    "adv1b": f"{SS_DIR}/adversarial_test1_slippery_slope_part2_1775157007592.png",
    "adv2a": f"{SS_DIR}/adversarial_test2_screenwriter_part1_1775157061831.png",
    "adv2b": f"{SS_DIR}/adversarial_test2_screenwriter_part2_1775157066636.png",
    "adv3a": f"{SS_DIR}/adversarial_test3_slang_part1_1775157122228.png",
    "adv3b": f"{SS_DIR}/adversarial_test3_slang_part2_1775157127604.png",
}

# ── Generated chart paths ─────────────────────────────────────────────────────
DATA_DIR = "data"
CHARTS = {
    "dataset":    f"{DATA_DIR}/fig_dataset_composition.png",
    "ml_eval":    f"{DATA_DIR}/fig_ml_evaluation.png",
    "approaches": f"{DATA_DIR}/fig_approach_comparison.png",
    "clusters":   f"{DATA_DIR}/cluster_visualization.png",
}

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(31, 73, 125)
    return p

def body(text, bold=False, italic=False, size=9.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 100, 100)

def add_table(headers, rows, font_size=9):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Shading Accent 1"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(font_size)
    for row_data in rows:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for run in row.cells[i].paragraphs[0].runs:
                run.font.size = Pt(font_size)
    return tbl

def insert_img(path_key, source="chart", width=5.5):
    path = CHARTS.get(path_key) if source == "chart" else SCREENSHOTS.get(path_key)
    if path and os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        body(f"[Image not found: {path_key}]", italic=True)

def bullet(title_text, desc_text):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(title_text + " ")
    r1.bold = True
    r1.font.size = Pt(10)
    p.add_run(desc_text).font.size = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_heading("SubstanceWatch AI", 0)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title_p.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 73, 125)

for text, size, bold in [
    ("Challenge 1: AI for Substance Abuse Risk Detection from Social Signals", 12, True),
    ("NSF NRT Research-A-Thon 2026  |  UMKC  |  Track A: AI Modeling & Reasoning", 10, False),
    ("Team: Joe Doan · Manan Koradiya · Ruixuan Hou · Aditya Naredla", 10, False),
    ("Submission: April 6, 2026  |  Demo: April 10, 2026  |  UMKC Student Union SU 401", 9, False),
]:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(size)
        run.bold = bold

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# §1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Introduction & Problem Statement")
body(
    "Substance abuse causes over 100,000 overdose deaths annually in the U.S. "
    "Warning signals frequently appear in online communities before escalating to physical crises, "
    "yet traditional public health monitoring relies on retrospective clinical reporting and delayed surveys. "
    "This project presents SubstanceWatch AI: an end-to-end AI pipeline that identifies substance abuse "
    "and emotional distress signals from anonymized social media data. Our core innovation is a "
    "Dual-Threshold Architecture combining a high-Recall ML Tripwire with a Gemini 2.5 LLM Evaluator "
    "backed by Retrieval-Augmented Generation (RAG), transforming noisy text into interpretable, "
    "evidence-grounded public health insights."
)


# ══════════════════════════════════════════════════════════════════════════════
# §2 — DATASET
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Dataset & Data Strategy")
body("We curated an ensemble of three public, anonymized datasets:")
add_table(
    ["Dataset", "Source / Size", "Purpose"],
    [
        ["Reddit Mental Health Classification", "HuggingFace · 1.1M posts",
         "Core: addiction, alcoholism, depression, PTSD, suicidewatch, anxiety"],
        ["Mental Health Corpus", "Kaggle · 28K posts",
         "Binary baseline: normal vs. emotional distress"],
        ["KUC Hackathon 2018 (reference)", "Kaggle · 200K reviews",
         "Pre-approved medical drug review reference"],
    ]
)
doc.add_paragraph()
body(
    "Stratified Sub-sampling: We engineered a targeted 43,551-row subset — ALL addiction (~7.7k) "
    "and alcoholism (~5.5k) posts; 5,000 from each of depression, suicidewatch, anxiety, ptsd, "
    "mentalhealth; and 5,000 safe control posts — ensuring dense, well-separated clusters for "
    "direct substance signals and emotional distress detection.",
    italic=True
)
doc.add_paragraph()
insert_img("dataset", source="chart", width=5.8)
caption("Figure 1: Stratified dataset composition — 43,551 posts across 10 subreddit categories. "
        "Red bars = High Risk classes; Green = Safe/Control.")


# ══════════════════════════════════════════════════════════════════════════════
# §3 — ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("3. System Architecture: Dual-Threshold Pipeline")
add_table(
    ["Layer", "Technology", "Role & Threshold"],
    [
        ["1. ML Tripwire", "TF-IDF + Logistic Regression",
         "Maximize Recall — cast a wide net. Threshold: 0.4. Output: 🟡 INITIAL SIGNAL DETECTED"],
        ["2. RAG Retrieval", "all-MiniLM-L6-v2 + ChromaDB (43k vectors)",
         "Retrieve high-confidence semantic context. Strict L2 Distance < 1.2 blocks hallucination."],
        ["3. LLM Evaluator", "Google Gemini 2.5 Flash",
         "Temporal + intent reasoning: Active Risk vs. Retrospective/Recovery. Output: 🔴/🟢 FINAL AI ASSESSMENT"],
    ]
)
doc.add_paragraph()
body(
    "Design Rationale: In public health, missing a genuine distress signal (False Negative) is more "
    "dangerous than a False Positive. The ML layer is tuned for maximum Recall (97.29%), while the "
    "LLM applies Precision-focused contextual reasoning to filter noise — correctly classifying "
    "retrospective/recovery narratives and promotional content as low risk.",
    italic=True
)
doc.add_paragraph()
insert_img("home", source="screenshot", width=5.5)
caption("Figure 2: SubstanceWatch AI Dashboard — Tab 1: Single Post Analysis; Tab 2: Batch CSV Analysis.")
doc.add_paragraph()
insert_img("full", source="screenshot", width=5.5)
caption("Figure 3: Dual-Threshold output — ML Tripwire flags the input; Gemini 2.5 confirms "
        "🔴 FINAL AI ASSESSMENT: HIGH RISK with cited Supporting Evidence.")


# ══════════════════════════════════════════════════════════════════════════════
# §4 — EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Evaluation Results")

heading("4.1  ML Baseline (80/20 Train-Test Split, 43,551 posts)", level=2)
body("The TF-IDF + Logistic Regression classifier was evaluated on an 80/20 split:")
add_table(
    ["Class", "Precision", "Recall", "F1 Score", "Support"],
    [
        ["Safe (Control)", "0.82", "0.93", "0.87", "1,002"],
        ["High Risk",      "0.9905", "0.9729", "0.9816", "7,709"],
        ["Weighted Avg",   "0.97",   "0.97",   "0.97",   "8,711"],
    ]
)
doc.add_paragraph()
insert_img("ml_eval", source="chart", width=5.8)
caption("Figure 4: ML layer evaluation — Left: Precision/Recall/F1 by class. "
        "Right: Overall 97% performance across all weighted metrics.")

doc.add_paragraph()
heading("4.2  Multi-Approach Comparison (60-Sample Balanced Test Set)", level=2)
body(
    "To validate our Dual-Threshold architecture, we ran all three pipeline components independently "
    "on a balanced held-out test set (30 High Risk + 30 Safe posts) and compared performance:"
)
add_table(
    ["Approach", "Precision", "Recall", "F1", "Accuracy"],
    [
        ["A — ML Only (TF-IDF + LR, threshold=0.4)", "0.9655", "0.9333", "0.9492", "0.9500"],
        ["B — RAG Only (ChromaDB, distance < 1.2)",   "0.9091", "1.0000", "0.9524", "0.9500"],
        ["C — Dual-Threshold (ML → RAG → LLM)",       "1.0000", "0.8333", "0.9091", "0.9167"],
    ]
)
doc.add_paragraph()
insert_img("approaches", source="chart", width=5.8)
caption("Figure 5: Multi-approach performance comparison. Each method makes a distinct "
        "precision-recall tradeoff — the Full Pipeline achieves Precision = 1.0 (zero false positives).")
doc.add_paragraph()
body(
    "Key Insight: The Full Pipeline (C) achieves perfect Precision (1.0) — zero false positives — "
    "by using LLM reasoning to correctly identify retrospective, fictional, and recovery narratives "
    "that would otherwise be false alarms. The RAG-only approach achieves perfect Recall (1.0) but "
    "at the cost of more false positives. Our Dual-Threshold design deliberately combines both: "
    "ML casts the wide net, LLM applies the precision filter.",
    italic=True
)


# ══════════════════════════════════════════════════════════════════════════════
# §5 — BEHAVIORAL PATTERN ANALYSIS (Core Task 2)
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Behavioral Pattern Analysis — Embedding Clustering")
body(
    "To satisfy Core Task 2 (Temporal and Behavioral Analysis), we applied K-Means clustering "
    "(k=8) on the 384-dimensional embedding vectors of 5,000 sampled High Risk posts stored in "
    "ChromaDB. Top TF-IDF keywords were extracted per cluster to automatically label emerging themes:"
)
add_table(
    ["Cluster", "Auto-Theme", "Size", "Dominant Subreddits"],
    [
        ["C0", "PTSD · Trauma · Flashbacks",             "574 (11.5%)", "ptsd 78%, mentalhealth 13%"],
        ["C1", "Depression · Hopelessness · Isolation",  "946 (18.9%)", "depression 33%, mentalhealth 25%, suicidewatch 22%"],
        ["C2", "Addiction · Weed · Drug Use",            "654 (13.1%)", "addiction 91%, alcoholism 4%"],
        ["C3", "Distress · Loneliness · Social Pain",    "688 (13.8%)", "depression 25%, mentalhealth 19%, suicidewatch 15%"],
        ["C4", "Drinking · Alcohol · Sobriety",          "635 (12.7%)", "alcoholism 85%, addiction 11%"],
        ["C5", "Anxiety · Panic · Work Stress",          "668 (13.4%)", "anxiety 65%, mentalhealth 18%"],
        ["C6", "Help-Seeking · Family Alcohol Concern",  "378 (7.6%)",  "addiction 43%, alcoholism 40%"],
        ["C7", "Suicidal Ideation · Crisis",             "457 (9.1%)",  "suicidewatch 67%, depression 19%"],
    ]
)
doc.add_paragraph()
insert_img("clusters", source="chart", width=5.8)
caption("Figure 6: PCA 2D visualization of 8 behavioral clusters from High Risk posts. "
        "Each color = one behavioral theme. Tight, well-separated clusters confirm the "
        "embedding space captures meaningful semantic distinctions.")
doc.add_paragraph()
body(
    "Finding: The clustering reveals that 'High Risk' is not monolithic — it comprises distinct "
    "behavioral profiles: active substance use (C2, C4), underlying emotional crises (C0, C1, C5), "
    "help-seeking on behalf of others (C6), and acute crisis/suicidal ideation (C7). "
    "This granularity enables targeted public health intervention strategies beyond binary classification.",
    italic=True
)


# ══════════════════════════════════════════════════════════════════════════════
# §6 — ADVERSARIAL TESTING
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Adversarial Testing (Stress-Test)")
body(
    "To validate reasoning beyond keyword matching, we designed three adversarial scenarios "
    "targeting known failure modes of naive classifiers. All three passed with expected outcomes:"
)
add_table(
    ["Scenario", "Trap Type", "ML Layer", "LLM Final", "Result"],
    [
        ["The Slippery Slope",
         "Recovery framing ('5 years clean') hiding imminent relapse — user parked outside old neighborhood, rationalizing 'just one hit'",
         "🟡 0.93", "🔴 HIGH RISK", "✅ PASS"],
        ["The Screenwriter",
         "Drug keywords ('opioid','cravings','dealer') inside fictional screenplay dialogue, attributed to a character not the author",
         "🟡 0.71", "🟢 LOW RISK",  "✅ PASS"],
        ["The Slang & Metaphor",
         "Zero explicit drug words — 'hitting the slopes / white noise / skiing' = cocaine slang, combined with job loss & panic",
         "🟡 0.93", "🔴 HIGH RISK", "✅ PASS"],
    ],
    font_size=8.5
)
doc.add_paragraph()
body(
    "Key finding: The Full Pipeline correctly handled all three scenarios. The ML layer flagged all inputs "
    "as potential signals (high Recall). The Gemini 2.5 LLM then applied contextual reasoning — "
    "correctly de-escalating the fictional screenwriter case (Precision filter) while escalating "
    "the recovery narrative and cocaine slang despite the absence of explicit drug keywords.",
    italic=True
)


# ══════════════════════════════════════════════════════════════════════════════
# §7 — INNOVATION HIGHLIGHTS & CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Innovation Highlights & Conclusion")
for title_text, desc in [
    ("Dual-Threshold Architecture:",
     "Separates high-Recall 'catching' (ML) from high-Precision 'filtering' (LLM+RAG) — "
     "a deliberate design for public health safety systems where false negatives are costly."),
    ("Multi-Approach Comparison:",
     "Quantitatively compared ML Only, RAG Only, and Full Pipeline on a 60-sample balanced test set. "
     "Full Pipeline achieves Precision = 1.0 (zero false positives)."),
    ("Behavioral Cluster Analysis:",
     "K-Means clustering on ChromaDB embeddings revealed 8 distinct behavioral risk profiles "
     "(PTSD, Suicidal Crisis, Active Substance Use, Help-Seeking, etc.) — enabling targeted intervention."),
    ("Active vs. Retrospective Classification:",
     "Gemini 2.5 differentiates recovery narratives, creative/fictional content, and slang-based "
     "active crisis posts — nuances no binary classifier can capture."),
    ("Grounded, Hallucination-Proof LLM Outputs:",
     "All LLM summaries are evidence-grounded via ChromaDB RAG. Strict L2 < 1.2 threshold "
     "blocks generation when no confident match exists."),
    ("Ethical AI by Design:",
     "No PII at any stage. Population-level analysis only. Data fully anonymized."),
]:
    bullet(title_text, desc)

doc.add_paragraph()
body(
    "Future Work: Integrate CDC/NIDA temporal datasets for population-level spike detection. "
    "Add BERTopic for emerging theme discovery. Deploy a temporal dashboard for community-level "
    "behavioral shift analysis by week/month."
)


# ══════════════════════════════════════════════════════════════════════════════
# §8 — REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("8. References")
refs = [
    "CDC National Drug Overdose Data: https://data.cdc.gov",
    "Reddit Mental Health Classification (HuggingFace): kamruzzaman-asif/reddit-mental-health-classification",
    "Mental Health Corpus (Kaggle): reihanenamdari/mental-health-corpus",
    "Reimers & Gurevych (2019). Sentence-BERT. arXiv:1908.10084",
    "Lewis et al. (2020). Retrieval-Augmented Generation. arXiv:2005.11401",
    "Google Gemini 2.5: https://deepmind.google/technologies/gemini/",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f"{i}. {ref}", style='List Number')
    for run in p.runs:
        run.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = "SubstanceWatch_AI_Report.docx"
doc.save(out)
print(f"✅ Report saved: {out}")
